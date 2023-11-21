import openai
import docx
from docx.shared import Pt

# OpenAI API key setup
openai.api_key = 'your-api-key'  # Replace with your actual OpenAI API key

# Function to generate text using OpenAI GPT
def generate_text(prompt, temperature=0.7, max_tokens=150):
    try:
        response = openai.completions.create(
            engine="text-davinci-004",  # Or another model version if you prefer
            prompt=prompt,
            temperature=temperature,
            max_tokens=max_tokens
        )
        return response.choices[0].text.strip()
    except openai.error.OpenAIError as e:
        print(f"An OpenAI error occurred: {e}")
        return None
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return None

# Create a new Word document
doc = docx.Document()
doc.add_heading('GPT-Generated Document', 0)

# Define the prompt for the GPT model
gpt_prompt = "Tell me a story about a lost kitten in the rain."

# Generate text using the function
generated_text = generate_text(gpt_prompt)

# Check if text was generated successfully
if generated_text:
    # Add generated text to the Word document
    p = doc.add_paragraph()
    run = p.add_run(generated_text)
    run.font.size = Pt(12)

    # Save the document
    doc.save('GPT_Generated_Document.docx')
    print("Document has been created and text has been written into it.")
else:
    print("Text generation was unsuccessful.")
